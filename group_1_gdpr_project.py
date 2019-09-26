"""
File name:  group_1_gdpr_project.py

Purpose of code:  CPCO will use information extracted from 5 data sources to create 5 dashboards, each with distinct information but that is all pertinent to data privacy regulations and the consequences of violating those regulations.  CPCO obtained its data by scraping JavaScript and html websites, and parsing a word file.  We also used an unusual source--Google Forms--to survey friends, family, and colleagues about their data privacy preferences.  We obtained 68 responses to 5 focused questions, each question offering the user the opportunity to select multiple answer choices or to provide additional comments.  We will use Python to delineate each of the unique answers that the users provided to each question and conduct calculations, including the percentage of users who selected each particular answer choice.

-  Dashboard A identifies specific US laws that meet the requirements of particular GDPR articles.

-  Dashboard B highlights European countries that have instituted privacy laws supplementing those mandated in the GDPR.  To view, please enter B into the input box
-  Dashboard C displays fines levied by the UK's Information Commissioner's Office (ICO) against British companies violating data privacy regulations, an example of one European country's domestic prosecutions citing GDPR.

-  Dashboard D demonstrates the most common methods hackers used in breaching data, the industries targeted most heavily, and the average number of records compromised per organization.

- Dashboard E provides a sample of customers' growing concerns about data privacy, based on a survey conducted by CP CO.

"""
#group_1_gdpr_project.py
# Carly Harris, Darshan Tina, Lokesh Nandiraju, Heather Manganello, Kopal Agrawal

#import statements
import textwrap
import pandas as pd
import csv
import datetime
import dash
import dash_core_components as dcc
import dash_html_components as html
import dash_table as dt
from dash.dependencies import Input, Output
import matplotlib.pyplot as plt
import plotly.graph_objs as go
import numpy as np
from decimal import Decimal
import io
from docx import Document
import re

#dashboard styling
external_stylesheets = ['https://codepen.io/chriddyp/pen/bWLwgP.css']
app = dash.Dash(__name__, external_stylesheets=external_stylesheets)

colors = {
    'background': '#D3E0F5'
}

# ************************** User Interface of Dash Board 1 ********************
#Read data from Word document into Dataframe
def read_docx(filename, tab_id=None, **kwargs):
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

#Parse raw data from .docx file
my_pan_file = read_docx('DS1_Uncleaned.docx')[0]

#Export read .docx content to a CSV file
my_pan_file.to_csv('DS_1_Cleaned.csv', sep=',')

####################################

#Read data from CSV file
my_panda = pd.read_csv('DS_1_Cleaned.csv')
del my_panda['Unnamed: 0']

### Take article number as input
print('\n\n##################################################################################################')
print('\nBeginning of execution on Data Source 1\n\n')

wanna_run_again=True
while wanna_run_again==True:
    art_num_good =False
    while art_num_good==False:
        try:
            art_no = int(input('\n\nPlease enter the article number: '))
            if 1<=art_no<=99:
                art_num_good=True
            else:
                print('Article number '+str(art_no)+' doesn\'t exist. Enter between 1 to 99 only')
        except:
            print('Bad format!. Please enter GDPR article number between 1 to 99 only')

    print('Article number '+str(art_no)+' is called \"'+str(my_panda.iloc[art_no,0])[3:]+'\"')

    ### Checking corresponding US LAW exists or not
    check_law_flag=False
    while check_law_flag==False:
        print()
        check_law=input('Do you want to check whether corresponding United States Law exists for the GDPR article number ' +str(art_no)+' ?\nEnter Y or N (Not case sensitive): ')
        if check_law=='Y' or check_law=='y':
            if str(my_panda.iloc[art_no,2])[0:3]=='N/A':
                print('\nNo corresponding U.S Law exists\nNON-COMPLIANT\n')
                check_law_flag=True
            else:
                #print Law exists and ask whether they want to see the text of that law
                print('\nRelevant U.S Law exists!!\nCOMPLIANT\n')
                check_display_law=False
                while check_display_law==False:
                    display_law=input('Do you want to view corresponding United States Law for GDPR article number ' +str(art_no)+' ?\nEnter Y or N (Not case sensitive): ')
                    if  display_law=='Y' or display_law=='y':
                        print('\nThe corresponding U.S Law for GDPR article number '+str(art_no)+' is: \n')
                        print('\"'+my_panda.iloc[art_no,2]+'\""')
                        check_law_flag=True
                        check_display_law=True
                    elif display_law=='N' or display_law=='n':
                        check_law_flag=True
                        check_display_law=True
                        print()
                    else:
                        print('Bad Input! Enter either Y or N: (Not case sensitive)')
        elif check_law=='N' or check_law=='n':
            print()
            check_law_flag=True
        else:
            print('Bad Input! Enter either Y or N (Not case sensitive): ')

    ### Ask user whether they want to see text for the GDPR article number they entered and display if they want to
    check_article_flag=False
    while check_article_flag==False:
        check_article=input('Do you want to view GDPR Regulation for article number '+str(art_no)+'?\n'+'Enter Y or N (Not case sensitive): ')
        if check_article=='Y' or check_article=='y':
            print()
            print('GDPR article number '+str(art_no)+' states that:')
            print()
            print('\"'+str(my_panda.iloc[art_no,1])+'\"')
            check_article_flag=True
            run_again_flag=False
            while run_again_flag==False:
                run_again=input('\nDo you want to run this again???\nEnter Y or N (Not case sensitive): \n')
                if run_again=='Y' or run_again=='y':
                    wanna_run_again=True
                    run_again_flag=True
                elif run_again=='N' or run_again=='n':
                    wanna_run_again=False
                    run_again_flag=True
                else:
                    run_again_flag=False
                    print('Bad input. Enter either Y or N (Not case sensitive): ')
        elif check_article=='N' or check_article=='n':
            check_article_flag=True
            run_again_flag=False
            while run_again_flag==False:
                run_again=input('\nDo you want to run this again???\nEnter Y or N (Not case sensitive): \n')
                if run_again=='Y' or run_again=='y':
                    wanna_run_again=True
                    run_again_flag=True
                elif run_again=='N' or run_again=='n':
                    wanna_run_again=False
                    run_again_flag=True
                else:
                    run_again_flag=False
                    print('Bad input. Enter either Y or N (Not case sensitive): ')

desc_lst=[]
for val in my_panda['GDPR & U.S Law Data.2']:
    if val[:3]=='N/A':
        desc_lst.append('Compliant')
    else:
        desc_lst.append('Not Compliant')
desc_pan=pd.Series(desc_lst)
print('\n\n\nCompliance Summary: \nThe number of GDPR articles for which compliant/corresponding U.S Laws exist are: '+str(desc_pan.describe()[3])+'.')
print('\nEnd of execution on Data Source 1')
print('##################################################################################################\n\n')


# ******************************* APPLICATION LAYOUT ********************************
# header layout
app.layout = html.Div(style={'backgroundColor': colors['background']}, children=[
    html.H1(children='CP CO',
            style={
                'textAlign': 'center'
                }),
    html.Div(children='GDPR and Information Security Consultancy',
             style={
                 'textAlign': 'center'
                 }),
    html.Div(children='Our mission: ',
             style={
                 'textAlign': 'center',
                 'font-weight': 'bold'
                 }),
    html.Div(children='To help organizations understand their vulnerabilities in data privacy protection and information security',
             style={
                 'textAlign': 'center'
                 }),
    html.Div(children='requirements for compliance with European regulations consequences for perceived violations, &',
             style={
                 'textAlign': 'center'
                 }),
    html.Div(children='avoid massive fines of up to 4% of revenue in order to effectively manage the use of European individuals personal information within their operations.',
             style={
                 'textAlign': 'center'
                 }),




    html.Div(children='Enter:'),
    html.Div(children='1-99 = Dashboard A: Correlation of US Laws to European Union\'s General Data Protection Regulation (GDPR) Articles '),
    html.Div(children='B = Dashboard B: Additional Data Privacy Requirements Per European Union Country'),
    html.Div(children='C = Dashboard C: Companies Fined by GDPR Law'),
    html.Div(children='D = Dashboard D: Statistics on Worldwide Data Breaches'),
    html.Div(children='E = Dashboard E: Surveys on Customers\' Preferences on Data Privacy'),
    html.Div(children=' '),
    html.Div(children='''
        Dashboard to Graph:
    '''),
    dcc.Input(id='input', value='', type='text'),
    html.Div(id='output-graph'),

])


#user input component
@app.callback(
    Output(component_id='output-graph', component_property='children'),
    [Input(component_id='input', component_property='value')]
)


#function to display graphs based on user input
def update_value(input_data):

    #-----------------------Dashboard A--------------------
    if(input_data.isdigit()):
        if(int(input_data) >= 0 and int(input_data) <= 99):

            # read cleaned csv data set
            df = pd.read_csv('DS_1_Cleaned.csv')
            df1 = pd.DataFrame(df)
            new = pd.DataFrame()
            new['Article'] = df1['GDPR & U.S Law Data']
            new['GDPR Law'] = df1['GDPR & U.S Law Data.1']
            new['US Law'] = df1['GDPR & U.S Law Data.2']

            temp = []
            vals = new.iloc[int(input_data)]
            t1 = pd.DataFrame(vals)
            for i in vals:
                temp.append(i)

            t1 = t1.transpose()
            # dash table of db1
            return dt.DataTable(
                    id='table',
                    columns=[{"name": i, "id": i} for i in new.columns],
                    data = t1.to_dict("rows"),
                    style_data={'whiteSpace': 'normal'},
                    css=[{
                        'selector': '.dash-cell div.dash-cell-value',
                        'rule': 'display: inline; white-space: inherit; overflow: inherit; text-overflow: inherit;'
                        }],
                    style_table={'overflowX': 'scroll'},

                    style_cell={
                        'textAlign': 'center',
                        'minWidth': '50px', 'maxWidth': '180px'
                        },
                    style_cell_conditional=[
                        {
                            'if': {'column_id': 'Region'},
                            'textAlign': 'center',
                            'if': {'row_index': 'odd'},
                            'backgroundColor': 'rgb(248, 248, 248)'
                            }
                        ],
                    style_header={
                        'backgroundColor': 'white',
                        'fontWeight':'bold'
                        }
                    )


    #-----------------------Dashboard B--------------------
    elif(input_data == 'B' or input_data == 'b'):

        # Web Scraping by implementing ParseHub API
        # http://www.alstongdprtracker.com/
        # Cleaning Raw Data Below
        """
        df = pd.read_csv('DataSet4_cleaned.csv')
        gdpr = pd.DataFrame()
        l1= []
        df = df.replace('Yes', 1)
        df = df.replace('No', 0)
        df = df.fillna(0)
        """
        # read csv DataSet and parse values
        df = pd.read_csv('DataSet4_cleaned.csv')
        gdpr = pd.DataFrame()
        df = df.replace('Yes', 1)
        df = df.replace('No', 0)
        df = df.fillna(0)
        gdpr = df.pivot(index='country', columns='article', values='deviation')
        #one_gdpr = np.arange(8)
        #one_gdpr = np.reshape(2,4)

        gdpr = gdpr.fillna(0)

        # get sum of supplemental laws for each country
        y_list = []
        for j in gdpr.values:
            sum = 0
            for i in j:
                sum += i
            y_list.append(sum)

        # get list of country whose article deviation value equals 1
        ls1 = []
        for i in gdpr.index:
            l1 = []
            for j in gdpr.columns:
                if(gdpr.loc[i][j] == 1.0):
                    l1.append(j)

            ls1.append(l1)

        # first trace value of bar graph
        trace1 = go.Bar(x=gdpr.index, y=y_list)

        # dictionary of gdpr article deviation per country
        valid = {}
        i=0
        for j in gdpr.index:
            valid[j] = ls1[i]
            i = i+1

        # fill empty values
        f1 = pd.DataFrame(dict([ (k,pd.Series(v)) for k,v in valid.items() ]))

        # matplotlib
        '''obj = gdpr.index
        y_pos = np.arange(len(obj))
        plt.bar(y_pos, ls1)
        plt.xticks(y_pos, obj)'''

        #return graphs - bar chart and data table
        return (dcc.Graph(
            id='example_graph',
            figure={
                'data':[trace1],
                'layout':
                go.Layout(title='Number of Supplemental Privacy Laws Enacted by European Union Country',
                          yaxis={'title':'Supplemental Privacy Laws'},
                          xaxis={'title':'Country'},
                                 barmode='stack', hovermode='closest')
                }),
                html.H4(children='Supplemental Laws to GDPR Articles per Country'),
                dt.DataTable(
                    id='table',
                    columns=[{"name": i, "id": i} for i in f1.columns],
                    data=f1.to_dict("rows"),
                    style_table={'overflowX': 'scroll'},
                    style_cell={
                        'textAlign': 'center',
                        },
                    style_cell_conditional=[
                        {
                            'if': {'column_id': 'Region'},
                            'textAlign': 'center',
                            'if': {'row_index': 'odd'},
                            'backgroundColor': 'rgb(248, 248, 248)'
                            }
                        ],
                    style_header={
                        'backgroundColor': 'white',
                        'fontWeight':'bold'
                        }
                    )

                )

    #-----------------------Dashboard C--------------------
    elif(input_data == 'C' or input_data == 'c'):

        # Web Scraping by implementing ParseHub API
        # https://ico.org.uk/action-weve-taken/enforcement/
        # Cleaning Raw Data Below
        """
        pd = pd.read_csv('C:\\Users\Darshan\Desktop\Dataset5_raw - Copy.csv')
        l = []
        l1 = []
        l = pd['Company_Fines']
        l1 = pd['Company_Dates']

        a = []
        for i in l1:
            a.append(i.split(',')[0])

        pd['Refined_Dates'] = a

        def isNumber(inputstr):
            if inputstr.isdigit():
                return True
            else:
                return False
        b = []
        for i in l:
            i=i.replace(' ','')
            s = "";
            if('£' in i):
                s = s+'£'
                j = 1
                while((i[i.find('£')+j])==',' or isNumber(i[i.find('£')+j])) :
                    s = s+i[i.find('£')+j]
                    j = j+1
                b.append(s)
            else:
                b.append(' ')

        pd['Fine_In_Euros'] = b
        New = pd.copy(deep = True)
        New = New.drop('Company_Fines',1)
        New = New.drop('Company_Dates',1)
        New.to_csv('DataSet5_cleaned.csv')
        New.to_json('test.json')
        """
        df = pd.read_csv('DataSet5_cleaned.csv')
        df.columns = ['Sr.no','Company_Name','Refined_Dates','Fine_In_Euros']

        company_name_list = df['Company_Name']
        date_list = df['Refined_Dates']
        fine_list = df['Fine_In_Euros']

        fine_in_numbers = []
        ind = 0
        index = []
        # get list of fines
        for i in fine_list:
            if(i!=' '):
                t = i[1:]
                x = t.find(',')
                t = t[0:x] + t[x+1:]
                fine_in_numbers.append(float(t))
                index.append(ind)
            ind+=1

        company_fined = []

        # specific companies fined
        for i in index:
            company_fined.append(company_name_list[i])

        # dict of companies and amount fined
        d = dict(zip(company_fined,fine_in_numbers))
        company_fined = []
        for i in d.keys():
            company_fined.append(i)
        fine_in_numbers = []
        for i in d.values():
            fine_in_numbers.append(i)

        # line graph
        return dcc.Graph(id='example',
                         figure ={
                             'data': [
                                 {'x':company_fined, 'y':fine_in_numbers, 'type':'line', 'name':'Fines Levied','xlabel':'Companies'},
                                 ],
                             'layout':{
                                 'title':'Companies Fined by GDPR Law',
                                 'xaxis':{
                                     'title':'Company',
                                     'showgrid':False,
                                     'zeroline':False,
                                     'tickmode':'array',
                                     'tickvals':[]
                                     },
                                 'yaxis':{
                                     'title':'Fine Levied'
                                     }
                                 }
                             })


    #-----------------------Dashboard D--------------------


    elif(input_data == 'D' or input_data == 'd'):

        """from bs4 import BeautifulSoup
        import requests
        website_data = requests.get('https://en.wikipedia.org/wiki/List_of_data_breaches').text
        soup = BeautifulSoup(website_data,'html.parser')
        table = soup.find('table',{'class':'wikitable sortable'})

        headers = [th.text for th in table.select("tr th")]

        with open("out_cleaned.csv", "w") as f:
            wr = csv.writer(f)
            wr.writerow(headers[:-1])
            wr.writerows([[td.text for td in row.find_all("td")][:-1] for row in table.select("tr + tr")])
            """


        #reading a csv file using pandas
        data=pd.read_csv('datasource6_cleaned.csv',index_col=False)
        for x in range(0,len(data['Method'])):
            #removing redundant values and characters
            data['Method'][x]=str(data['Method'][x]).replace('unknown\n','unknown')
            data['Method'][x]=str(data['Method'][x]).replace('nan','unknown')
            data['Method'][x]=str(data['Method'][x]).replace('poor security\n','poor security')
            data['Method'][x]=str(data['Method'][x]).replace('hacked\n','hacked')
            data['Method'][x]=str(data['Method'][x]).replace('social engineering\n','social engineering')
            data['Method'][x]=str(data['Method'][x]).replace('lost / stolen media\n','lost / stolen media')
        #extracting distinct Method values
        setMethod=set(data['Method'])
        setMethod=list(setMethod)
        listZero=[0,0,0,0,0,0,0,0,0,0]
        dictMethod=dict(zip(setMethod,listZero))
        for x in range(0,len(data['Method'])):
            #calculating frequency of each method
            for y in setMethod:
                if data['Method'][x]==y:
                    dictMethod[y]+=1

        key_list = []
        for i in dictMethod.keys():
            key_list.append(i)
        val_list = []
        for j in dictMethod.values():
            val_list.append(j)


        setOT=set(data['Organization type'])

        for x in range(0,len(data['Organization type'])):
            #removing redundant characters and values
            data['Organization type'][x]=str(data['Organization type'][x]).replace('web, tech','tech, web')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('consulting, accounting\n','consulting, accounting')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('social network\n','social network')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('web\n','web')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('ticket distribution\n','ticket distribution')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('government, database\n','government, database')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('transport\n','transport')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('genealogy\n','genealogy')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('healthcare\n','healthcare')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('financial\n','financial')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('credit reporting\n','credit reporting')
            data['Organization type'][x]=str(data['Organization type'][x]).replace('retail\n','retail')
        setOT=set(data['Organization type'])
        setOT=list(setOT)
        listZero=[0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
        dictOT=dict(zip(setOT,listZero))
        for x in range(0,len(data['Organization type'])):
            #calculating frequency of each Organiztion Type attacked
            for y in setOT:
                if data['Organization type'][x]==y:
                    dictOT[y]+=1

        k2 = []
        v2 = []
        for i in dictOT.keys():
            k2.append(i)
        for j in dictOT.values():
            v2.append(j)

        # data frame
        df = pd.read_csv('datasource6_cleaned.csv')
        n=[]
        j=0
        x = []
        #cleaning Records to only obtain integer values
        for i in df['Records']:
            j+=1
            if re.match('^[1-9][0-9]*',i):
                n.append(i)
                x.append(j)
        t=[]
        j=0
        for i in n:
            j+=1
            if re.search(r'^[^a-zA-Z]*$',i):
                t.append(i)
            else:
                x.remove(j)

        s = []
        for i in t:
            if '\n' in i:
                i = i[:i.find('\n')]
            s.append(i)
        q = []
        m=[]
        for i in s:
            q = i.split(',')
            string =""
            for i in q:
                string += i
            m.append(float(string))
        df.columns = ['Entity', 'Year', 'Records', 'Organization type', 'Method']
        year = []
        for i in df['Year']:
            if '\n' in i:
                i = i[:i.find('\n')]
            year.append(i)
        new_year =[]
        count =0
        for i in range(len(year)):
            if i not in x:
                year = year[:i]+year[i+1:]
        year = year[1:]
        d ={}

        for i in range(len(year)):
            if i not in d.keys():
                d[year[i]]=m[i]
            else:
                d[year[i]] += m[i]
        new_year = []
        new_values = []
        for i in d.keys():
            new_year.append(i)
        for i in d.values():
            new_values.append(i)

        a=np.array(df['Year'])
        b=np.array(df['Entity'])
        c=np.array([])
        for x in a:
            x.replace('\n','')

            c=np.append(c,[x])
        attack=pd.crosstab(b,c)
        attack[attack==1]='Compromised'
        e = []
        kopal = pd.DataFrame(attack)
        for i in kopal.index:
            e.append(i)
        kopal['Company_Names'] = e
        kopal = kopal.iloc[:,[16,0,1,2,3,4,5,6,7,8,9,10,11,12,13,14,15]]



        # bar graphs and data table
        return(html.Div([
    html.Div([
        html.Div([
            dcc.Graph(id='g1', figure={
                'data': [
                    {'x': key_list, 'y': val_list, 'type': 'bar'}],

                 'layout':
                 go.Layout(title='Number of Each Attack Vector Used',
                           yaxis={'title':'# of Attacks'},
                           barmode='stack', hovermode='closest')
                })
            ],
                 className="six columns"),


        html.Div([
            dcc.Graph(id='g2', figure={
                'data': [
                    {'x': k2, 'y': v2, 'type': 'bar'}],
                'layout':
                 go.Layout(title='Number of Attacks per Organization',
                           yaxis={'title':'# of Attacks'},
                           barmode='stack', hovermode='closest')
                })
        ], className="six columns"),
    ], className="row"),
    ]),
               dcc.Graph(
        id='example-graph',
        figure={
            'data': [
                {'x': new_year, 'y': new_values, 'type': 'bar'},
            ],
            'layout': {
                'title': 'Number of Records Compromised per Year',
                'xaxis':{
                    'title': 'Year'
                    },
                'yaxis':{
                    'title': '# of Records'
                    }
            }
        }
    ),
    html.H4(children='Years in which Companies Data were Compromised'),
                dt.DataTable(
                    id='table',
                    columns=[{"name": i, "id": i} for i in kopal.columns],
                    data=attack.to_dict("rows"),
                    style_table={'overflowX': 'scroll'},
                    style_cell={
                        'textAlign': 'center',
                        },
                    style_cell_conditional=[
                        {
                            'if': {'column_id': 'Region'},
                            'textAlign': 'center',
                            'if': {'row_index': 'odd'},
                            'backgroundColor': 'rgb(248, 248, 248)'
                            }
                        ],
                    style_header={
                        'backgroundColor': 'white',
                        'fontWeight':'bold'
                        }
                    )


)

    #-----------------------Dashboard E--------------------
    elif(input_data == 'E' or input_data == 'e'):

        pd1 = pd.read_csv('data_source3_raw_68.csv') #import raw data file, downloaded from Google Forms
        # Google Forms outoupts its information in a csv file, not requiring any scraping or parsing
        # This file contains responses from 68 individuals surveyed about their privacy preferences.
        # The form contained 5 questions, and respondents could select more than one answer to each.

        ShinyTbl = pd1.copy(deep=True)
        ShinyTbl = ShinyTbl.drop('Timestamp',1) #don't need the dates or times of the responses

        #ShinyTbl.describe()

        ShinyTbl.columns = ['Advertisements','Not_In_Favor','Feelings_PI','Companies_Know_You','GDPR_provisions'] #rename column headers

        # Note that respondents can select one or MORE of these options (so cells can contain several answers).
        # While most answers are distinct, searches for "no" responses would inaccurately also count "not sure" responses.
        # Therefore, need to adjust responses for the "Advertisements" and "Companies_Know_You" columns.

        for b in range(0,len(ShinyTbl['Advertisements'])):
            ShinyTbl['Advertisements'][b]=str(ShinyTbl['Advertisements'][b]).replace('Not sure','Unsure') #replace "not sure" with "unsure"

        for f in range(0,len(ShinyTbl['Companies_Know_You'])):
            ShinyTbl['Companies_Know_You'][f]=str(ShinyTbl['Companies_Know_You'][f]).replace('Not sure','Unsure')

        ShinyTbl.to_csv('DataSource_3_cleaned.csv') #outputing cleaned file

        # Counting Specific Responses to 1st question ("Advertisements")

        L1 = ShinyTbl['Advertisements']

        opt1_1 = 'protect'; opt2_1 = 'careful'; opt3_1 = 'No'; opt4_1 = 'Unsure'
        ctopt1_1 = 0; ctopt2_1 = 0; ctopt3_1 = 0; ctopt4_1 = 0

        for a in L1:
                a = (str)(a)
                if opt1_1 in a:
                    ctopt1_1 += 1
                if opt2_1 in a:
                    ctopt2_1 += 1
                if opt3_1 in a:
                    ctopt3_1 += 1
                if opt4_1 in a:
                    ctopt4_1 += 1
        totalAns1 = ctopt1_1 + ctopt2_1 + ctopt3_1 + ctopt4_1

        dict1 = {'Yes, even if they\'re not being careful with my data; I don\'t care': ctopt2_1,
                'Yes, as long as they protect and use my data responsibly' : ctopt1_1,
                'No' : ctopt3_1,
                'Not Sure' : ctopt4_1,
                'Total Responses' : totalAns1}

        # Counting Specific Responses to 2nd question ("Not In Favor")

        L2 = ShinyTbl['Not_In_Favor']

        opt1_2 = 'misusing'; opt2_2 = 'not protecting'; opt3_2 = 'watching'
        ctopt1_2 = 0; ctopt2_2 = 0; ctopt3_2 = 0

        for c in L2:
                c = (str)(c)
                if opt1_2 in c:
                    ctopt1_2 += 1
                if opt2_2 in c:
                    ctopt2_2 += 1
                if opt3_2 in c:
                    ctopt3_2 += 1

        totalAns2 = ctopt1_2 + ctopt2_2 + ctopt3_2

        dict2 = {'I am concerned about companies misusing my data' : ctopt1_2,
                'I am concerned about companies not protecting my data from hackers' : ctopt2_2,
                'I just don\'t like companies watching what I do online' : ctopt3_2,
                'Total Responses' : totalAns2}

        # Counting Specific Responses to 3rd question ("Feelings_PI")

        L3 = ShinyTbl['Feelings_PI']

        opt1_3 = 'mostly good'; opt2_3 = 'bad'; opt3_3 = 'care about'; opt4_3 = 'Not sure'
        ctopt1_3 = 0; ctopt2_3 = 0; ctopt3_3 = 0; ctopt4_3 = 0

        for d in L3:
                d = (str)(d)
                if opt1_3 in d:
                    ctopt1_3 += 1
                if opt2_3 in d:
                    ctopt2_3 += 1
                if opt3_3 in d:
                    ctopt3_3 += 1
                if opt4_3 in d:
                    ctopt4_3 += 1

        totalAns3 = ctopt1_3 + ctopt2_3 + ctopt3_3 + ctopt4_3

        dict3 = {'I think that they are mostly good at using our data in responsible and secure ways' : ctopt1_3,
                'I think that they mostly want to use our data in a responsible and secure way . . . but they are bad at it' : ctopt2_3,
                'I don\'t think that they care about using our data responsibly or securely at all' : ctopt3_3,
                'Not sure' : ctopt4_3,
                'Total Responses' : totalAns3}

        # Counting Specific Responses to 4th question ("Companies_Know_You")

        L4 = ShinyTbl['Companies_Know_You']

        opt1_4 = 'Yes'; opt2_4 = 'No'; opt3_4 = 'Unsure'
        ctopt1_4 = 0; ctopt2_4 = 0; ctopt3_4 = 0

        for e in L4:
                e = (str)(e)
                if opt1_4 in e:
                    ctopt1_4 += 1
                if opt2_4 in e:
                    ctopt2_4 += 1
                if opt3_4 in e:
                    ctopt3_4 += 1

        totalAns4 = ctopt1_4 + ctopt2_4 + ctopt3_4

        dict4 = {'Yes' : ctopt1_4,
                'No' : ctopt2_4,
                'Not sure' : ctopt3_4,
                'Total Responses' : totalAns4}

        # Counting Specific Responses to 5th question ("Feelings_PI")

        L5 = ShinyTbl['GDPR_provisions']

        opt1_5 = 'delete'; opt2_5 = 'being done'; opt3_5 = 'control'; opt4_5 = 'complain'
        ctopt1_5 = 0; ctopt2_5 = 0; ctopt3_5 = 0; ctopt4_5 = 0

        for g in L5:
                g = (str)(g)
                if opt1_5 in g:
                    ctopt1_5 += 1
                if opt2_5 in g:
                    ctopt2_5 += 1
                if opt3_5 in g:
                    ctopt3_5 += 1
                if opt4_5 in g:
                    ctopt4_5 += 1

        totalAns5 = ctopt1_5 + ctopt2_5 + ctopt3_5 + ctopt4_5

        dict5 = {'The right to ask companies to delete the data they have on me' : ctopt1_5,
                'The right to ask what is being done with my data' : ctopt2_5,
                'The ability to control how my data is being used' : ctopt3_5,
                'The ability to complain when my data is misused, which could result in steep fines for the offending company' : ctopt4_5,
                'Total Responses' : totalAns5}

        # question 1
        d1key = []
        d1Pct = []

        for i in dict1.keys():
            d1key.append(i)
        del d1key[-1]

        for i in dict1.values():
            d1Pct.append(i)

        total = d1Pct[-1]
        for i in range(4):
            d1Pct[i] = (d1Pct[i]/total) *100
            d1Pct[i] = round(d1Pct[i],2)

        del d1Pct[-1]

        #question 2
        d2key = []
        d2Pct = []
        for i in dict2.keys():
            d2key.append(i)
        del d2key[-1]

        for i in dict2.values():
            d2Pct.append(i)

        total = d2Pct[-1]

        for i in range(4):
            d2Pct[i] = (d2Pct[i]/total) *100
            d2Pct[i] = round(d2Pct[i],2)

        del d2Pct[-1]

        #question 3
        d3key = []
        d3Pct = []
        for i in dict3.keys():
            d3key.append(i)
        del d3key[-1]

        for i in dict3.values():
            d3Pct.append(i)

        total = d3Pct[-1]

        for i in range(4):
            d3Pct[i] = (d3Pct[i]/total) *100
            d3Pct[i] = round(d3Pct[i],2)

        del d3Pct[-1]

        #question 4
        d4key = []
        d4Pct = []

        for i in dict4.keys():
            d4key.append(i)
        del d4key[-1]

        for i in dict4.values():
            d4Pct.append(i)

        total = d4Pct[-1]
        for i in range(3):
            d4Pct[i] = (d4Pct[i]/total) *100
            d4Pct[i] = round(d4Pct[i],2)

        del d4Pct[-1]

        #question 5
        d5key = []
        d5Pct = []

        for i in dict5.keys():
            d5key.append(i)
        del d5key[-1]

        for i in dict5.values():
            d5Pct.append(i)

        total = d5Pct[-1]

        for i in range(4):
            d5Pct[i] = (d5Pct[i]/total) *100
            d5Pct[i] = round(d5Pct[i],2)

        del d5Pct[-1]

        q1 = 'Are you in favor of websites or apps using what they learn about you to show you advertisements that you might find interesting?'
        q1 = textwrap.wrap(q1)
        q2 = 'If you are not in favor of receiving targeted advertisements, why not?'
        q2 = textwrap.wrap(q2)
        q3 = 'In general, what are your feelings toward companies that host websites today and have access to our personal information?'
        q3 = textwrap.wrap(q3)
        q4 = 'Do you think companies hosting websites know too much about you?'
        q4 = textwrap.wrap(q4)
        q5 = 'Which of these types of privacy provisions--which already have been enacted in Europe--would you most like see applied in the US?'
        q5 = textwrap.wrap(q5)

        return (
            dcc.Graph(id='pi',
                         figure={
                             'data': [
                                 {
                                     'labels': d1key,
                                     'values': d1Pct,
                                     'type': 'pie',
                                     'title': '<br>'.join(q1),
                                     'marker': {'colors': ['rgb(114,119,121)',
                                                           'rgb(18, 36, 37)',
                                                           'rgb(34, 53, 101)',
                                                           'rgb(51,155,255)',
                                                           'rgb(6, 4, 4)']},
                                     'domain': {'x': [0, .48],
                                                'y': [0, .49]},
                                     'hoverinfo':'label+percent+name',
                                     'textinfo':'none'
                                     },
                                 {
                                     'labels': d2key,
                                     'values': d2Pct,
                                     'title': '<br>'.join(q2),
                                     'marker': {'colors': ['rgb(114,119,121)',
                                                           'rgb(18, 36, 37)',
                                                           'rgb(34, 53, 101)',
                                                           'rgb(51,155,255)',
                                                           'rgb(6, 4, 4)']},
                                     'type': 'pie',

                                     'domain': {'x': [.52, 1],
                                                'y': [0, .49]},
                                     'hoverinfo':'label+percent+name',
                                     'textinfo':'none'
                                     },
                                 {
                                     'labels': d3key,
                                     'values': d3Pct,
                                     'title':'<br>'.join(q3),
                                     'marker': {'colors': ['rgb(114,119,121)',
                                                           'rgb(18, 36, 37)',
                                                           'rgb(34, 53, 101)',
                                                           'rgb(51,155,255)',
                                                           'rgb(6, 4, 4)']},
                                     'type': 'pie',

                                     'domain': {'x': [0, .33],
                                                'y': [.5, 1]},
                                     'hoverinfo':'label+percent+name',
                                     'textinfo':'none'
                                     },
                                 {
                                     'labels': d4key,
                                     'values': d4Pct,
                                     'title': '<br>'.join(q4),
                                     'marker': {'colors': ['rgb(114,119,121)',
                                                           'rgb(18, 36, 37)',
                                                           'rgb(34, 53, 101)',
                                                           'rgb(51,155,255)',
                                                           'rgb(6, 4, 4)']},
                                     'type': 'pie',

                                     'domain': {'x': [.35, .67],
                                                'y': [.51, 1]},
                                     'hoverinfo':'label+percent+name',
                                     'textinfo':'none'
                                     },
                                 {
                                     'labels': d5key,
                                     'values': d5Pct,
                                     'title': '<br>'.join(q5),
                                     'marker': {'colors': ['rgb(114,119,121)',
                                                           'rgb(18, 36, 37)',
                                                           'rgb(34, 53, 101)',
                                                           'rgb(51,155,255)',
                                                           'rgb(6, 4, 4)']},
                                     'type': 'pie',

                                     'domain': {'x': [.70, 1],
                                                'y': [.51, 1]},
                                     'hoverinfo':'label+percent+name',
                                     'textinfo':'none'
                                     }
                                 ],
                             'layout': {'title': 'Customers\' Preferences on Data Privacy',
                                        'height':700,
                                        'showlegend': False}
                             }
                         )
            )

    #------------------Invalid User input------------------
    else:
        return html.Div('Please enter: 1-99 or B,C,D,E')







#run server on main module
if __name__ == '__main__':
    app.run_server(debug=False)
